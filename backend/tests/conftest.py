"""Programmatic test fixtures — no external files needed.

All fixtures are generated in memory (Pillow+piexif for images,
pypdf for PDFs, python-docx+lxml for DOCX) so the suite is fully
self-contained.
"""
from __future__ import annotations

from io import BytesIO

import piexif
import pytest
from PIL import Image, PngImagePlugin


# ── GPS helper ────────────────────────────────────────────────────────────────

def _dms(degrees: int, minutes: int, seconds_tenths: int) -> list[tuple[int, int]]:
    """Return DMS as piexif rational triples.  seconds_tenths is seconds × 10."""
    return [(degrees, 1), (minutes, 1), (seconds_tenths, 10)]


# ── Fixtures ──────────────────────────────────────────────────────────────────

@pytest.fixture(scope="session")
def jpeg_with_gps() -> bytes:
    """100×100 JPEG carrying GPS coords (San Francisco), camera info, and author."""
    img = Image.new("RGB", (100, 100), (73, 109, 137))

    exif_dict: dict = {
        "0th": {
            piexif.ImageIFD.Make: b"Apple",
            piexif.ImageIFD.Model: b"iPhone 14 Pro",
            piexif.ImageIFD.Software: b"17.4.1",
            piexif.ImageIFD.Artist: b"Jane Doe",
            piexif.ImageIFD.DateTime: b"2023:08:15 10:30:00",
        },
        "Exif": {
            piexif.ExifIFD.DateTimeOriginal: b"2023:08:15 10:30:00",
            piexif.ExifIFD.DateTimeDigitized: b"2023:08:15 10:30:01",
        },
        "GPS": {
            piexif.GPSIFD.GPSLatitudeRef: b"N",
            piexif.GPSIFD.GPSLatitude: _dms(37, 46, 296),   # 37°46'29.6"N
            piexif.GPSIFD.GPSLongitudeRef: b"W",
            piexif.GPSIFD.GPSLongitude: _dms(122, 25, 98),  # 122°25'9.8"W
            piexif.GPSIFD.GPSAltitudeRef: 0,
            piexif.GPSIFD.GPSAltitude: (52, 1),
        },
        "1st": {},
    }
    exif_bytes = piexif.dump(exif_dict)

    buf = BytesIO()
    img.save(buf, format="JPEG", exif=exif_bytes)
    return buf.getvalue()


@pytest.fixture(scope="session")
def jpeg_no_exif() -> bytes:
    """Plain 100×100 JPEG with no EXIF data at all."""
    img = Image.new("RGB", (100, 100), (200, 200, 200))
    buf = BytesIO()
    img.save(buf, format="JPEG")
    return buf.getvalue()


@pytest.fixture(scope="session")
def jpeg_no_gps_but_camera() -> bytes:
    """JPEG with camera metadata but no GPS."""
    img = Image.new("RGB", (80, 80), (50, 80, 120))
    exif_dict: dict = {
        "0th": {
            piexif.ImageIFD.Make: b"Canon",
            piexif.ImageIFD.Model: b"EOS R5",
            piexif.ImageIFD.Software: b"Firmware 1.8.0",
        },
        "Exif": {
            piexif.ExifIFD.DateTimeOriginal: b"2024:01:20 14:00:00",
        },
        "GPS": {},
    }
    buf = BytesIO()
    img.save(buf, format="JPEG", exif=piexif.dump(exif_dict))
    return buf.getvalue()


@pytest.fixture(scope="session")
def png_with_text() -> bytes:
    """PNG carrying tEXt metadata chunks (Author, Software, Comment)."""
    img = Image.new("RGB", (60, 60), (200, 100, 50))
    meta = PngImagePlugin.PngInfo()
    meta.add_text("Author", "Bob Smith")
    meta.add_text("Software", "GIMP 2.10.36")
    meta.add_text("Comment", "Test fixture — do not distribute")
    buf = BytesIO()
    img.save(buf, format="PNG", pnginfo=meta)
    return buf.getvalue()


@pytest.fixture(scope="session")
def corrupt_bytes() -> bytes:
    """Bytes that start like a JPEG but are otherwise garbage."""
    return b"\xff\xd8\xff\xe0" + b"\x00" * 16 + b"not valid jpeg content at all"


# ── PDF fixtures ──────────────────────────────────────────────────────────────

@pytest.fixture(scope="session")
def pdf_with_metadata() -> bytes:
    """Single-page PDF with /Author, /Creator, /Title, and /CreationDate set."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.add_metadata({
        "/Author": "Alice Example",
        "/Creator": "Microsoft Word 16.0",
        "/Title": "Confidential Report",
        "/Subject": "Test fixture",
        "/Keywords": "test metadata fixture",
        "/CreationDate": "D:20240315120000+00'00'",
        "/ModDate": "D:20240315130000+00'00'",
    })
    buf = BytesIO()
    writer.write(buf)
    return buf.getvalue()


@pytest.fixture(scope="session")
def pdf_no_metadata() -> bytes:
    """Single-page PDF with no /Info dictionary."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buf = BytesIO()
    writer.write(buf)
    return buf.getvalue()


# ── DOCX fixtures ─────────────────────────────────────────────────────────────

@pytest.fixture(scope="session")
def docx_with_metadata() -> bytes:
    """DOCX with core properties set and one tracked insertion injected via lxml."""
    from docx import Document
    from lxml import etree

    doc = Document()
    props = doc.core_properties
    props.author = "Bob Builder"
    props.last_modified_by = "Carol Editor"
    props.title = "Project Proposal"
    props.subject = "Testing"
    props.revision = 5

    # Add content so paragraphs[0] exists, then inject a tracked insertion
    doc.add_paragraph("Base content.")
    _W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W = f"{{{_W_NS}}}"
    para = doc.paragraphs[0]._element
    ins = etree.SubElement(para, f"{W}ins")
    ins.set(f"{W}id", "1")
    ins.set(f"{W}author", "Bob Builder")
    ins.set(f"{W}date", "2024-03-15T12:00:00Z")
    run = etree.SubElement(ins, f"{W}r")
    t = etree.SubElement(run, f"{W}t")
    t.text = "tracked inserted text"

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture(scope="session")
def docx_no_metadata() -> bytes:
    """Blank DOCX with no properties set."""
    from docx import Document

    doc = Document()
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
